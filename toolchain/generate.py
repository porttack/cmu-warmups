"""generate.py — build Word workbooks from content.csv.

One warm-up = exactly two pages (page 1: header + START HERE strip + vocabulary;
page 2: Part 1 + Part 2), plus a unit cover page. All look-and-feel lives in the
THEME dict below. Figures come from the `figure` spec column via figrender.py.

    python generate.py                 # every unit + full course doc
    python generate.py --unit 1        # just unit 1
    python generate.py --course cs1 --out out
"""
import argparse
import csv
import os
import re
import tempfile

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import figrender

THEME = {
    "font": "Calibri",
    "mono": "Consolas",
    "base_pt": 12,
    "small_pt": 9,
    "strip_pt": 10,
    "title_pt": 13,
    "accent": "1F3A5F",      # title bar / rules
    "shade": "EEEEEE",       # strip letter cells, Part 2 bar
    "cover_tag": "1F3A5F",
    "line_h_pt": 22,          # writing-line height (student handwriting room)
    "line_color": "99AABB",
    "code_bg": "F5F5F5",
    "code_border": "888888",
    "margin_tb": 0.55,
    "margin_lr": 0.7,
    "fig_dpi": 150,           # px width -> inches divisor is 96; render at 1500px
}

# Fallback check-ins for row E. Only consulted when a page has no meta/checkin
# row of its own — an authored check-in is always used verbatim.
SEL = [
    "One word for how today is going so far:",
    "Something you are looking forward to this week:",
    "A person you could ask for help if you get stuck:",
    "One thing that is going well outside of this class:",
    "If you finish early today, what will you try next?",
    "A goal — small is fine — for this class period:",
]


# ----------------------------- low-level xml helpers -----------------------------
def _shade(el, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    el.append(shd)


def _p_border(p, edges=("bottom",), color="000000", sz=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    for e in edges:
        b = OxmlElement("w:" + e)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        pbdr.append(b)


def _cell_shade(cell, hexcolor):
    _shade(cell._tc.get_or_add_tcPr(), hexcolor)


def _no_space(p, before=0, after=0, line_pt=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_pt is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_pt)


def _set_table_full(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    w = OxmlElement("w:tblW")
    w.set(qn("w:type"), "pct")
    w.set(qn("w:w"), "5000")  # 100%
    tblPr.append(w)


def _run(p, text, size=None, bold=False, color=None, mono=False, italic=False):
    r = p.add_run(text)
    r.font.name = THEME["mono"] if mono else THEME["font"]
    if size:
        r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


# Inline markdown: **bold**, *italic*, `code`. Mirrors mdInline()/mdRuns() in
# wblib.js/wbdocx.js so all three renderers treat the same content the same way.
_MD_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _md_run(p, text, size=None, bold=False, color=None):
    text = "" if text is None else str(text)
    last = 0
    wrote = False
    for m in _MD_RE.finditer(text):
        if m.start() > last:
            _run(p, text[last:m.start()], size=size, bold=bold, color=color)
        tok = m.group(0)
        if tok.startswith("**"):
            _run(p, tok[2:-2], size=size, bold=True, color=color)
        elif tok.startswith("*"):
            _run(p, tok[1:-1], size=size, bold=bold, color=color, italic=True)
        else:
            _run(p, tok[1:-1], size=size, bold=bold, color=color, mono=True)
        last = m.end()
        wrote = True
    if last < len(text):
        _run(p, text[last:], size=size, bold=bold, color=color)
        wrote = True
    if not wrote:
        _run(p, text, size=size, bold=bold, color=color)


# ----------------------------- content pieces -----------------------------
def writing_lines(doc, n, line_pt=None):
    line_pt = line_pt or THEME["line_h_pt"]
    for _ in range(int(n)):
        p = doc.add_paragraph()
        _no_space(p, line_pt=line_pt)
        _p_border(p, ("bottom",), THEME["line_color"], sz=6)


def title_bar(doc, w):
    p = doc.add_paragraph()
    _no_space(p, before=2, after=6)
    _shade(p._p.get_or_add_pPr(), THEME["accent"])
    _run(p, "Warm-up %s " % w["no"], size=THEME["title_pt"], bold=True, color="FFFFFF")
    _run(p, " —  %s" % w["topic"], size=THEME["title_pt"], color="FFFFFF")


def header_block(doc, w):
    t = doc.add_table(rows=2, cols=3)
    t.style = "Table Grid"
    _set_table_full(t)
    cells = [("Name", "Date", "Period"), ("Partner", "Points ___ / ___", "Marked by:  ☐ me   ☐ partner")]
    for r in range(2):
        for c in range(3):
            cell = t.cell(r, c)
            cell.width = Inches([3.6, 1.9, 1.5][c])
            # label bottom-aligned with writing room above it (space before, not after)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            p = cell.paragraphs[0]
            _no_space(p, before=14)
            _run(p, cells[r][c], size=THEME["small_pt"])
    title_bar(doc, w)


def strip_block(doc, w):
    lab = doc.add_paragraph()
    _no_space(lab, before=8, after=2)
    _run(lab, "START HERE", size=THEME["strip_pt"], bold=True)

    ican = w["ican"] if w["ican"] else ["I can start today's warm-up.", "I can ask for help when stuck."]
    no = int(w["no"]) if str(w["no"]).isdigit() else 0
    # (letter, label, opts, has_line): opts (circle-one words) are flushed right; None = plain row
    items = [
        ("A", "Pace check — circle one:", "feeling good    a bit distracted    stuck — I could use help", False),
        ("B", (ican[0] if len(ican) > 0 else ""), "got it    shaky    not yet", False),
        ("C", (ican[1] if len(ican) > 1 else ""), "got it    shaky    not yet", False),
        ("D", w.get("reflect") or "Looking back at last class — one thing that clicked, and one thing still fuzzy:", None, True),
        ("E", w.get("checkin") or SEL[no % len(SEL)], None, True),
    ]
    t = doc.add_table(rows=len(items), cols=2)
    t.style = "Table Grid"
    _set_table_full(t)
    for i, (letter, label, opts, has_line) in enumerate(items):
        lc = t.cell(i, 0)
        lc.width = Inches(0.3)
        _cell_shade(lc, THEME["shade"])
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _no_space(lp)
        _run(lp, letter, size=THEME["strip_pt"], bold=True)
        rc = t.cell(i, 1)
        rc.width = Inches(6.7)
        rp = rc.paragraphs[0]
        _no_space(rp, after=2)
        if opts is not None:
            rp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
            _run(rp, label + "\t" + opts, size=THEME["strip_pt"])
        else:
            _run(rp, label, size=THEME["strip_pt"])
        if has_line:
            lnp = rc.add_paragraph()
            _no_space(lnp, line_pt=THEME["line_h_pt"])
            _p_border(lnp, ("bottom",), THEME["line_color"], sz=6)


def section_label(doc, text):
    p = doc.add_paragraph()
    _no_space(p, before=10, after=3)
    r = _run(p, text, size=THEME["base_pt"], bold=True)
    _p_border(p, ("bottom",), THEME["accent"], sz=12)


def part2_bar(doc):
    p = doc.add_paragraph()
    _no_space(p, before=10, after=4)
    _shade(p._p.get_or_add_pPr(), THEME["shade"])
    _run(p, "PART 2 ", bold=True, size=THEME["base_pt"])
    _run(p, "— keep going if you finish early.", size=THEME["base_pt"], color="444444")


def code_box(doc, text):
    p = doc.add_paragraph()
    _no_space(p, before=4, after=4)
    _shade(p._p.get_or_add_pPr(), THEME["code_bg"])
    _p_border(p, ("top", "bottom", "left", "right"), THEME["code_border"], sz=4)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        _run(p, line, mono=True, size=11)


def hint_n(hint, default):
    for tok in str(hint or "").replace(" ", "").split(";"):
        if tok.startswith("n="):
            try:
                return int(tok[2:])
            except ValueError:
                pass
    return default


def hint_w(hint, default):
    for tok in str(hint or "").replace(" ", "").split(";"):
        if tok.startswith("w="):
            try:
                return int(tok[2:])
            except ValueError:
                pass
    return default


def place_figure(doc, spec, wpx, figcache):
    if spec not in figcache:
        fd, path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        figrender.render_to(spec, path, px=1500)
        figcache[spec] = path
    p = doc.add_paragraph()
    _no_space(p, before=4, after=4)
    run = p.add_run()
    run.add_picture(figcache[spec], width=Inches(min(6.5, wpx / 96.0)))


def render_item(doc, it, figcache):
    t = it["type"]
    if t == "figure":
        place_figure(doc, it["figure"] or "grid", hint_w(it["hint"], 230), figcache)
    elif t == "code":
        code_box(doc, it["content"])
    elif t == "label":
        p = doc.add_paragraph()
        _no_space(p, before=6, after=2)
        _md_run(p, it["content"], bold=True, size=THEME["base_pt"])
    elif t == "lines":
        writing_lines(doc, hint_n(it["hint"], 3))
    elif t == "vocab":
        p = doc.add_paragraph()
        _no_space(p, before=4, after=2)
        _md_run(p, it["content"], bold=True, size=THEME["base_pt"])
        writing_lines(doc, hint_n(it["hint"], 2))
    else:  # p
        p = doc.add_paragraph()
        _no_space(p, before=4, after=2)
        _md_run(p, it["content"], size=THEME["base_pt"])
        n = hint_n(it["hint"], 2)
        if n > 0:
            writing_lines(doc, n)


# ----------------------------- model -----------------------------
def load_rows(path):
    with open(path, newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def warmup_from_rows(rows):
    meta = rows[0]
    w = {"course": meta["course"], "unit": meta["unit"], "no": meta["warmup_no"],
         "page": meta["page"], "topic": "", "ican": [], "reflect": "", "checkin": "",
         "vocab": [], "part1": [], "part2": []}
    for r in rows:
        sec = r["section"]
        it = {"type": r["type"], "content": r["content"], "hint": r["hint"], "figure": r["figure"]}
        if sec == "meta":
            if r["type"] == "topic":
                w["topic"] = r["content"]
            elif r["type"] == "ican":
                w["ican"].append(r["content"])
            elif r["type"] == "reflect":
                w["reflect"] = r["content"]
            elif r["type"] == "checkin":
                w["checkin"] = r["content"]
        elif sec in ("vocab", "part1", "part2"):
            w[sec].append(it)
    return w


def unit_pages(rows, course, unit):
    seen, order = set(), []
    for r in rows:
        if r["course"] == course and str(r["unit"]) == str(unit) and r["page"].startswith("w"):
            if r["page"] not in seen:
                seen.add(r["page"])
                order.append(r["page"])
    return order


# ----------------------------- document assembly -----------------------------
def new_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = THEME["font"]
    st.font.size = Pt(THEME["base_pt"])
    for s in doc.sections:
        s.top_margin = Inches(THEME["margin_tb"])
        s.bottom_margin = Inches(THEME["margin_tb"])
        s.left_margin = Inches(THEME["margin_lr"])
        s.right_margin = Inches(THEME["margin_lr"])
    return doc


def cover_page(doc, course, unit, ws):
    p = doc.add_paragraph()
    _no_space(p, after=2)
    _run(p, course.upper(), bold=True, color=THEME["cover_tag"], size=12)
    h = doc.add_paragraph()
    _no_space(h, after=2)
    _run(h, "Unit %s" % unit, bold=True, size=34)
    vocab_ct = sum(len(w["vocab"]) for w in ws)
    sub = doc.add_paragraph()
    _no_space(sub, after=12)
    _run(sub, "%d warm-ups  \u00b7  %d vocabulary terms" % (len(ws), vocab_ct), color="555555", size=12)

    box = doc.add_paragraph()
    _no_space(box, before=6, after=2)
    _run(box, "By the end of this unit, I can\u2026", bold=True, size=THEME["base_pt"])
    seen = set()
    for w in ws:
        for s in w["ican"]:
            if s and s not in seen:
                seen.add(s)
                li = doc.add_paragraph(style="List Bullet")
                _no_space(li, after=1)
                _run(li, "\u2610 " + s, size=11)

    hb = doc.add_paragraph()
    _no_space(hb, before=12, after=2)
    _run(hb, "How to use these pages", bold=True, size=THEME["base_pt"])
    ht = doc.add_paragraph()
    _no_space(ht, after=2)
    _run(ht, "Each warm-up is two pages. Page 1 is the header, the START HERE strip (A\u2013E), and "
             "Vocabulary. Page 2 is Part 1 (everyone) and Part 2 (keep going if you finish early). "
             "Mark your own or a partner\u2019s work \u2014 it does not count against you.", size=11)


def build_warmup(doc, w):
    header_block(doc, w)               # page 1
    strip_block(doc, w)
    figcache = doc._figcache
    if w["vocab"]:                     # no vocab rows -> no heading, no gap
        section_label(doc, "Vocabulary \u2014 write each in your own words")
        for it in w["vocab"]:
            render_item(doc, it, figcache)
    doc.add_page_break()               # -> page 2
    section_label(doc, "Part 1 \u2014 core work")
    for it in w["part1"]:
        render_item(doc, it, figcache)
    part2_bar(doc)
    for it in w["part2"]:
        render_item(doc, it, figcache)


def build_unit(rows, course, unit, out_dir):
    doc = new_doc()
    doc._figcache = {}
    ws = [warmup_from_rows([r for r in rows if r["page"] == pg and r["course"] == course and str(r["unit"]) == str(unit)])
          for pg in unit_pages(rows, course, unit)]
    cover_page(doc, course, unit, ws)
    for w in ws:
        doc.add_page_break()
        build_warmup(doc, w)
    name = "%s-Unit%s-Workbook.docx" % (course.upper(), unit)
    path = os.path.join(out_dir, name)
    doc.save(path)
    return path, len(ws)


def build_course(rows, course, out_dir):
    doc = new_doc()
    doc._figcache = {}
    unitset = sorted({r["unit"] for r in rows if r["course"] == course}, key=lambda u: int(u))
    first = True
    for unit in unitset:
        ws = [warmup_from_rows([r for r in rows if r["page"] == pg and r["course"] == course and str(r["unit"]) == str(unit)])
              for pg in unit_pages(rows, course, unit)]
        if not first:
            doc.add_page_break()
        first = False
        cover_page(doc, course, unit, ws)
        for w in ws:
            doc.add_page_break()
            build_warmup(doc, w)
    path = os.path.join(out_dir, "%s-Full-Workbook.docx" % course.upper())
    doc.save(path)
    return path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--csv", default="content.csv")
    ap.add_argument("--course", default="cs1")
    ap.add_argument("--unit", default=None)
    ap.add_argument("--out", default="out")
    a = ap.parse_args()
    os.makedirs(a.out, exist_ok=True)
    rows = load_rows(a.csv)
    if a.unit:
        path, n = build_unit(rows, a.course, a.unit, a.out)
        print("built", path, "(%d warm-ups)" % n)
    else:
        units = sorted({r["unit"] for r in rows if r["course"] == a.course}, key=lambda u: int(u))
        for u in units:
            path, n = build_unit(rows, a.course, u, a.out)
            print("built", path, "(%d warm-ups)" % n)
        print("built", build_course(rows, a.course, a.out))


if __name__ == "__main__":
    main()
